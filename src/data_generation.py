#!/usr/bin/env python3
"""
generate_10000_strict_docs.py

Generates DOCX documents strictly following user constraints.
Set NUM_DOCS at top (default 10000). Test with small value first.
Produces: docs/*.docx and manifest.csv
"""

import os, random, math, shutil, textwrap
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
from tqdm import tqdm

# ---------------- CONFIG ----------------
NUM_DOCS = 10000        # change to 20 for quick test; set to 10000 for full run
OUTPUT_DIR = "docs"
TMP_IMG_DIR = "tmp_imgs"
MANIFEST_CSV = "manifest.csv"

# Quota minimums (enforced)
IMAGE_PERCENT = 0.30    # >=30%
TABLE_PERCENT = 0.25    # >=25%
EQUATION_PERCENT = 0.18 # >=18%
COMPLEX_PERCENT = 0.60  # >=60%

# Page distribution parameters
MEAN_PAGES = 3.0
SD_PAGES = 1.5
MIN_PAGES = 1
MAJORITY_MAX = 10
TAIL_COUNT_15_30 = 200
OUTLIERS_30_PLUS = 50

# Words per page approximation (tune if you find pages too short)
WORDS_PER_PAGE = 380

# Max paragraphs per section (controls density)
MAX_PARAS_PER_SECTION = 6

# Topics (diverse)
TOPICS = {
    "STEM": ["Quantum Entanglement", "Numerical Methods for PDEs", "Neural Network Optimization", "Applied Thermodynamics"],
    "Humanities": ["Renaissance Art", "Existentialism in Literature", "History of the Silk Road", "Comparative Philosophy"],
    "SocialSciences": ["Behavioral Economics", "Social Network Analysis", "Psychological Resilience", "Urban Sociology"],
    "CurrentEvents": ["Renewable Energy Policies", "Global Supply Chain Vulnerabilities", "AI Ethics and Regulation"],
    "Reference": ["How-to: Setting up a Local Dev Environment", "RESTful API Design", "Guide: Data Cleaning in Python"]
}
# ----------------------------------------

# create dirs
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TMP_IMG_DIR, exist_ok=True)
random.seed(42)
np.random.seed(42)

# ---------- helpers ----------
def sample_page_counts(n):
    """Create page count list satisfying tail and outlier requirements."""
    if n < (TAIL_COUNT_15_30 + OUTLIERS_30_PLUS):
        # scale tails down if n too small
        tail = max(0, n // 50)
        out = max(0, n // 200)
    else:
        tail = TAIL_COUNT_15_30
        out = OUTLIERS_30_PLUS
    majority = n - (tail + out)
    counts = []
    for _ in range(majority):
        val = int(round(max(MIN_PAGES, min(MAJORITY_MAX, random.gauss(MEAN_PAGES, SD_PAGES)))))
        counts.append(val)
    for _ in range(tail):
        counts.append(random.randint(15, 30))
    for _ in range(out):
        counts.append(random.randint(30, 60))
    # fill/trim
    while len(counts) < n:
        counts.append(int(round(max(MIN_PAGES, min(MAJORITY_MAX, random.gauss(MEAN_PAGES, SD_PAGES))))))
    if len(counts) > n:
        counts = counts[:n]
    random.shuffle(counts)
    return counts

def make_paragraph_text(topic, min_sent=4, max_sent=10):
    templates = [
        f"{topic} has been a subject of extensive study, combining theoretical and empirical approaches.",
        f"Recent work on {topic.lower()} emphasizes reproducibility and rigorous evaluation.",
        f"In many contexts, {topic.lower()} interacts with policy, technological advances, and social outcomes.",
        f"Researchers approach {topic.lower()} using quantitative analysis, case studies, and simulation.",
        f"Important challenges around {topic.lower()} include scalability, robustness, and ethical considerations."
    ]
    s_count = random.randint(min_sent, max_sent)
    return " ".join(random.choices(templates, k=s_count))

def save_chart(path, title="Chart"):
    x = np.linspace(0, 10, 140)
    y = np.sin(x) + np.random.normal(scale=0.2, size=x.shape)
    plt.figure(figsize=(6,3))
    plt.plot(x, y)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(path, dpi=150)
    plt.close()

def save_equation_image(path, text="E = mc^2"):
    img = Image.new("RGB", (700,100), color="white")
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 26)
    except:
        font = ImageFont.load_default()
    d.text((10,10), text, font=font, fill=(0,0,0))
    img.save(path)

def add_table_docx(doc, rows=4, cols=3):
    table = doc.add_table(rows=rows+1, cols=cols)
    # header
    for j in range(cols):
        table.rows[0].cells[j].text = f"Header {j+1}"
    for i in range(1, rows+1):
        for j in range(cols):
            table.rows[i].cells[j].text = str(round(random.uniform(0,100),2))

# ---------- prepare quotas ----------
def prepare_quotas(n):
    num_images = math.ceil(IMAGE_PERCENT * n)
    num_tables = math.ceil(TABLE_PERCENT * n)
    num_eq = math.ceil(EQUATION_PERCENT * n)
    num_complex = math.ceil(COMPLEX_PERCENT * n)
    indices = list(range(n))
    random.shuffle(indices)
    image_idxs = set(indices[:num_images])
    table_idxs = set(indices[num_images:num_images+num_tables])
    eq_idxs = set(indices[num_images+num_tables:num_images+num_tables+num_eq])
    complex_idxs = set(indices[num_images+num_tables+num_eq:
                             num_images+num_tables+num_eq+num_complex])
    return image_idxs, table_idxs, eq_idxs, complex_idxs

# ---------- main generation ----------
def generate_all(num_docs=NUM_DOCS):
    page_counts = sample_page_counts(num_docs)
    image_idxs, table_idxs, eq_idxs, complex_idxs = prepare_quotas(num_docs)
    manifest = []

    print(f"Generating {num_docs} docs into '{OUTPUT_DIR}'. THIS MAY TAKE AWHILE.")
    for i in tqdm(range(num_docs), desc="Generating docs"):
        idx = i + 1
        pages = page_counts[i]
        topic_group = random.choice(list(TOPICS.keys()))
        topic = random.choice(TOPICS[topic_group])
        title = f"{topic} — Analysis Report"

        doc = Document()
        # set normal style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # Title only (NO generated/author lines)
        h = doc.add_heading(title, level=0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        total_words_target = pages * WORDS_PER_PAGE
        words_written = 0
        section_count = max(3, min(12, pages))

        has_image = False
        has_table = False
        has_equation = False
        is_complex = False

        for s in range(section_count):
            lvl = random.choice([1,2,2,3])
            sec_title = f"Section {s+1}: {random.choice(['Overview','Background','Methods','Results','Discussion','Applications'])}"
            doc.add_heading(sec_title, level=lvl)

            # paragraphs
            paras_in_section = random.randint(1, MAX_PARAS_PER_SECTION)
            for _p in range(paras_in_section):
                if words_written >= total_words_target:
                    break
                para_text = make_paragraph_text(topic, min_sent=4, max_sent=12)
                # sometimes extend paragraph to increase length
                if random.random() < 0.35:
                    para_text += " " + make_paragraph_text(topic, min_sent=2, max_sent=6)
                doc.add_paragraph(para_text)
                words_written += len(para_text.split())

            # complex formatting
            if (i in complex_idxs) and (not is_complex) and random.random() < 0.9:
                is_complex = True
                # bullet list
                for li in range(random.randint(2,5)):
                    doc.add_paragraph(make_paragraph_text(topic, min_sent=1, max_sent=3)[:200], style='List Bullet')
                # numbered list sometimes
                if random.random() < 0.5:
                    for li in range(random.randint(2,4)):
                        doc.add_paragraph(f"{li+1}. " + make_paragraph_text(topic, min_sent=1, max_sent=3)[:200], style='List Number')

            # insert table if forced and not already added
            if (i in table_idxs) and (not has_table) and random.random() < 0.6:
                rows = random.randint(3, 7)
                cols = random.randint(2, 5)
                add_table_docx(doc, rows=rows, cols=cols)
                has_table = True

            # insert equation image if forced and not already added
            if (i in eq_idxs) and (not has_equation) and random.random() < 0.45:
                eq_file = os.path.join(TMP_IMG_DIR, f"eq_{idx}_{s}.png")
                eq_text = random.choice(["E = mc^2", "∫_0^∞ e^{-x} dx = 1", "f(x)=ax^2+bx+c", "∇·E = ρ/ε0"])
                save_equation_image(eq_file, eq_text)
                try:
                    doc.add_picture(eq_file, width=Inches(3))
                    has_equation = True
                    # remove temp image file to save disk space
                    try:
                        os.remove(eq_file)
                    except:
                        pass
                except:
                    pass

            # insert chart image if forced and not already added
            if (i in image_idxs) and (not has_image) and random.random() < 0.6:
                img_file = os.path.join(TMP_IMG_DIR, f"chart_{idx}_{s}.png")
                save_chart(img_file, title=f"{topic} chart")
                try:
                    doc.add_picture(img_file, width=Inches(5))
                    has_image = True
                    # remove temp image after embedding
                    try:
                        os.remove(img_file)
                    except:
                        pass
                except:
                    pass

            if words_written >= total_words_target:
                break

        # References / citations section (inline-style, not external)
        doc.add_heading("References", level=2)
        for r in range(random.randint(2,6)):
            doc.add_paragraph(f"[{r+1}] Example reference on {topic}.")

        # Save docx
        fname = f"doc_{idx:06d}.docx"
        docx_path = os.path.join(OUTPUT_DIR, fname)
        try:
            doc.save(docx_path)
        except Exception as e:
            print("Failed to save", docx_path, ":", e)
            continue

        manifest.append({
            "filename": fname,
            "pages_target": pages,
            "topic": topic,
            "has_image": int(has_image),
            "has_table": int(has_table),
            "has_equation": int(has_equation),
            "is_complex": int(is_complex),
            "path": docx_path
        })

    # ----- post-check: ensure quotas satisfied; if not, insert missing items into random docs -----
    df = pd.DataFrame(manifest)
    current_images = int(df['has_image'].sum())
    current_tables = int(df['has_table'].sum())
    current_eq = int(df['has_equation'].sum())
    current_complex = int(df['is_complex'].sum())

    needed_images = max(0, math.ceil(IMAGE_PERCENT * num_docs) - current_images) if (num_docs := NUM_DOCS) else 0
    needed_tables = max(0, math.ceil(TABLE_PERCENT * num_docs) - current_tables)
    needed_eq = max(0, math.ceil(EQUATION_PERCENT * num_docs) - current_eq)
    needed_complex = max(0, math.ceil(COMPLEX_PERCENT * num_docs) - current_complex)

    # Helper to re-open and insert into docx
    def insert_image_to_doc(path):
        try:
            doc = Document(path)
            img_file = os.path.join(TMP_IMG_DIR, f"fix_img_{random.randint(100000,999999)}.png")
            save_chart(img_file, title="Inserted chart")
            doc.add_paragraph()
            doc.add_picture(img_file, width=Inches(5))
            doc.save(path)
            os.remove(img_file)
            return True
        except:
            return False

    def insert_table_to_doc(path):
        try:
            doc = Document(path)
            doc.add_paragraph()
            add_table_docx(doc, rows=4, cols=3)
            doc.save(path)
            return True
        except:
            return False

    def insert_eq_to_doc(path):
        try:
            doc = Document(path)
            eq_file = os.path.join(TMP_IMG_DIR, f"fix_eq_{random.randint(100000,999999)}.png")
            save_equation_image(eq_file, "E = mc^2")
            doc.add_paragraph()
            doc.add_picture(eq_file, width=Inches(3))
            doc.save(path)
            os.remove(eq_file)
            return True
        except:
            return False

    # Do fixes on random subset if necessary
    if needed_images > 0:
        candidates = df[df['has_image'] == 0].sample(min(needed_images, len(df[df['has_image'] == 0])))
        for p in candidates['path']:
            if insert_image_to_doc(p):
                df.loc[df['path'] == p, 'has_image'] = 1

    if needed_tables > 0:
        candidates = df[df['has_table'] == 0].sample(min(needed_tables, len(df[df['has_table'] == 0])))
        for p in candidates['path']:
            if insert_table_to_doc(p):
                df.loc[df['path'] == p, 'has_table'] = 1

    if needed_eq > 0:
        candidates = df[df['has_equation'] == 0].sample(min(needed_eq, len(df[df['has_equation'] == 0])))
        for p in candidates['path']:
            if insert_eq_to_doc(p):
                df.loc[df['path'] == p, 'has_equation'] = 1

    # For complex formatting, insert lists in random files lacking them
    if needed_complex > 0:
        candidates = df[df['is_complex'] == 0].sample(min(needed_complex, len(df[df['is_complex'] == 0])))
        for path in candidates['path']:
            try:
                doc = Document(path)
                doc.add_paragraph("Complex formatting added: ", style='List Bullet')
                doc.add_paragraph("1. Added numbered item", style='List Number')
                doc.save(path)
                df.loc[df['path'] == path, 'is_complex'] = 1
            except:
                pass

    # Save final manifest
    df.to_csv(MANIFEST_CSV, index=False)
    print("DONE: Generated docs:", len(df))
    print("Manifest:", MANIFEST_CSV)
    return df

# Run generator
if __name__ == "__main__":
    # Quick sanity check
    if NUM_DOCS > 2000:
        print("WARNING: NUM_DOCS is large. Ensure you have enough disk space before running.")
        print("You can test with NUM_DOCS = 20 first.")
    df = generate_all(NUM_DOCS)
    print(df[['has_image','has_table','has_equation','is_complex']].sum())
